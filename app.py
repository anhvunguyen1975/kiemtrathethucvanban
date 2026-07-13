import streamlit as st
from docx import Document
from docx.shared import Cm, Pt
from docx.oxml.ns import qn  
from docx.shared import Pt, RGBColor
from docx.oxml import OxmlElement
import re
import io
import os
import unicodedata  
import json  

#from reportlab.pdfbase import pdfmetrics
#from reportlab.pdfbase.ttfonts import TTFont
#from reportlab.pdfgen import canvas
#from reportlab.lib.pagesizes import letter

# --- CẤU HÌNH GIAO DIỆN STREAMLIT ---
st.set_page_config(page_title="Kiểm tra thể thức văn bản NĐ 30", page_icon="💧", layout="wide", initial_sidebar_state="expanded")
custom_css = """
<style>
  .stAppDeployButton {display: none !important;}
  #MainMenu { visibility: hidden !important; }
  header {visibility: hidden;}
  footer {visibility: hidden;}
  .viewerBadge_container__171of {display: none !important;}
  
  [data-testid="collapsedControl"] {
      display: flex !important;
      visibility: visible !important;
      z-index: 999999 !important; 
      background-color: transparent !important;
  }
</style>
"""
st.markdown(custom_css, unsafe_allow_html=True)

# --- DANH MỤC LOẠI VĂN BẢN CHUẨN NĐ 30 ---
LOAI_VAN_BAN_CHUAN = {
    "NQ": "Nghị quyết", "QĐ": "Quyết định", "CT": "Chỉ thị",
    "QC": "Quy chế", "QyĐ": "Quy định", "TB": "Thông báo", "HD": "Hướng dẫn", 
    "CTr": "Chương trình", "CTY": "Công ty", "KH": "Kế hoạch", "PA": "Phương án", 
    "BC": "Báo cáo", "BB": "Biên bản", "TTr": "Tờ trình", "HĐ": "Hợp đồng", 
    "GUQ": "Giấy ủy quyền", "GM": "Giấy mời", "GGT": "Giấy giới thiệu"
}
MAP_CHUAN_HOA_LOAI = {k.upper(): k for k in LOAI_VAN_BAN_CHUAN.keys()}
MAP_CHUAN_HOA_LOAI.update({"TTR": "TTr", "CTR": "CTr", "QYĐ": "QyĐ"})

CONFIG_FILE = "bawa_config.json"
def load_hidden_config():
    if not os.path.exists(CONFIG_FILE):
        default_config = {
            "ten_co_quan_me": "CÔNG TY CỔ PHẦN CẤP NƯỚC BẠC LIÊU",
            "dia_danh_chuan": "Bạc Liêu",
            "danh_sach_phong_ban": [
                {"ten": "Phòng Tổ chức Hành chính", "viet_tat": "TCHC"},
                {"ten": "Phòng Kế toán", "viet_tat": "KT"},
                {"ten": "Phòng Kinh doanh", "viet_tat": "KD"},
                {"ten": "Phòng Kế hoạch Kỹ thuật", "viet_tat": "KHKT"},
                {"ten": "Ban kiểm soát", "viet_tat": "BKS"},
                {"ten": "Xí nghiệp Cấp nước", "viet_tat": "XNCN"}
            ]
        }
        with open(CONFIG_FILE, "w", encoding="utf-8") as f:
            json.dump(default_config, f, ensure_ascii=False, indent=4)
        return default_config
    with open(CONFIG_FILE, "r", encoding="utf-8") as f:
        return json.load(f)

HIDDEN_CONFIG = load_hidden_config()

def init_vietnamese_fonts():
    win_font_reg, win_font_bold = "C:\\Windows\\Fonts\\arial.ttf", "C:\\Windows\\Fonts\\arialbd.ttf"
    if os.path.exists(win_font_reg) and os.path.exists(win_font_bold):
        pdfmetrics.registerFont(TTFont('Arial', win_font_reg))
        pdfmetrics.registerFont(TTFont('Arial-Bold', win_font_bold))
        return 'Arial', 'Arial-Bold'
    return 'Helvetica', 'Helvetica-Bold'

PDF_FONT_REG, PDF_FONT_BOLD = init_vietnamese_fonts()

def set_font_times(run):
    run.font.name = 'Times New Roman'
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.get_or_add_rFonts()
    rFonts.set(qn('w:ascii'), 'Times New Roman')
    rFonts.set(qn('w:hAnsi'), 'Times New Roman')
    rFonts.set(qn('w:eastAsia'), 'Times New Roman')
    rFonts.set(qn('w:cs'), 'Times New Roman')

# --- HÀM KIỂM TRA ĐƯỜNG KẺ NGANG (CHỈ CẢNH BÁO) ---
def check_agency_line_comprehensive(doc):
    warnings = []
    return warnings

# --- HÀM KIỂM TRA SÂU ---
def analyze_document_v6(doc):
    success_items = []
    error_list = []
    warning_list = []
    ambiguous_dict = {} 

    for idx, section in enumerate(doc.sections):
        top_cm = round(section.top_margin.cm, 2) if section.top_margin else 0
        bottom_cm = round(section.bottom_margin.cm, 2) if section.bottom_margin else 0
        left_cm = round(section.left_margin.cm, 2) if section.left_margin else 0
        right_cm = round(section.right_margin.cm, 2) if section.right_margin else 0
        
        margin_errors = []
        if not (2.0 <= top_cm <= 2.5): margin_errors.append(f"Lề trên: {top_cm}cm")
        if not (2.0 <= bottom_cm <= 2.5): margin_errors.append(f"Lề dưới: {bottom_cm}cm")
        if not (3.0 <= left_cm <= 3.5): margin_errors.append(f"Lề trái: {left_cm}cm")
        if not (1.5 <= right_cm <= 2.0): margin_errors.append(f"Lề phải: {right_cm}cm")
        
        if margin_errors:
            error_list.append(f"❌ **Sai căn lề trang:** " + " | ".join(margin_errors))
        else:
            success_items.append(f"Kích thước căn lề trang Section {idx+1} (Đạt chuẩn)")

    wrong_fonts_detected = set()
    all_paragraphs = []
    for p in doc.paragraphs:
        if p.text.strip(): all_paragraphs.append(p)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    if paragraph.text.strip(): all_paragraphs.append(paragraph)

    for p in all_paragraphs:
        for run in p.runs:
            if run.text.strip() and run.font.name and run.font.name != "Times New Roman":
                wrong_fonts_detected.add(run.font.name)

    if wrong_fonts_detected:
        error_list.append(f"❌ **Sai Phông chữ:** Phát hiện phông lạ {list(wrong_fonts_detected)} → Bắt buộc dùng **Times New Roman**.")
    else:
        success_items.append("Toàn bộ phông chữ bản thảo (Đạt chuẩn: Times New Roman)")

    def is_paragraph_bold(p):
        style_bold = False
        try:
            if p.style and p.style.font and p.style.font.bold: style_bold = True
        except: pass
        text_runs = [run for run in p.runs if run.text.strip()]
        if not text_runs: return False
        for run in text_runs:
            rb = run.bold if run.bold is not None else style_bold
            if not rb: return False
        return True

    # --- KIỂM TRA ĐỊNH DẠNG TÊN CƠ QUAN (CHỈ QUÉT VÙNG HEADER) ---
    agency_checked = False
    elements_to_check = []
    if doc.tables:
        for row in doc.tables[0].rows:
            for cell in row.cells:
                elements_to_check.extend(cell.paragraphs)
    elements_to_check.extend(doc.paragraphs[:5])

    for p in elements_to_check:
        text_clean = p.text.replace("|", "").strip()
        if not text_clean: continue
        text_lower = text_clean.lower()
        if "cấp nước bạc liêu" in text_lower or text_lower == "công ty cổ phần":
            if text_clean != text_clean.upper() or is_paragraph_bold(p):
                error_list.append(f"❌ - [Lỗi] Tên cơ quan chủ quản: `[{text_clean}]` tại góc trái văn bản chưa được VIẾT HOA hoặc in đậm.")
            agency_checked = True
            break 

    doc_text = "\n".join([p.text for p in all_paragraphs])
    has_noi_nhan = False
    has_nguoi_ky = False
    chuc_vu_keywords = ["GIÁM ĐỐC", "PHÓ GIÁM ĐỐC", "CHỦ TỊCH", "TRƯỞNG PHÒNG", "PHÒ TRƯỞNG PHÒNG", "PHÓ PHÒNG", "TỔNG GIÁM ĐỐC", "PHÓ TỔNG GIÁM ĐỐC"]

    for p in all_paragraphs:
        line_clean = p.text.replace("|", "").strip()
        if not line_clean: continue
        line_upper = unicodedata.normalize('NFC', line_clean.upper())
        line_clean_spaces = re.sub(r'\s+', ' ', line_clean)
        
        if "CỘNG HÒA XÃ HỘI CHỦ NGHĨA VIỆT NAM" in line_upper:
            if not is_paragraph_bold(p): error_list.append(f"❌ Sai Quốc hiệu: `[{line_clean}]` → Bắt buộc phải **IN ĐẬM**.")
        if "độc lập" in line_clean.lower() and "hạnh phúc" in line_clean.lower():
            if not is_paragraph_bold(p): error_list.append(f"❌ Sai Tiêu ngữ: `[{line_clean}]` → Bắt buộc phải **IN ĐẬM**.")

        if any(kv in line_clean_spaces for kv in chuc_vu_keywords):
            if "KÍNH GỬI" not in line_upper:
                has_nguoi_ky = True

        if line_clean.lower().startswith("kính gửi"):
            if "KÍNH GỬI" in line_clean:
                error_list.append(f"❌ Sai định dạng Kính gửi: `[{line_clean[:20]}...]` → Theo NĐ 30, chữ 'Kính gửi' phải in thường, không viết HOA toàn bộ.")
            if is_paragraph_bold(p):
                error_list.append(f"❌ Sai kiểu chữ Kính gửi: `[{line_clean[:20]}...]` → Bắt buộc dùng kiểu chữ thường, đứng (không in đậm).")

        if line_clean.lower().startswith("nơi nhận"):
            has_noi_nhan = True
            is_italic = any(r.italic for r in p.runs if r.text.strip())
            is_bold = is_paragraph_bold(p)
            if not (is_bold and is_italic):
                error_list.append(f"❌ Khối `Nơi nhận:` chưa đúng định dạng (Yêu cầu: Vừa **In đậm** vừa *In nghiêng*).")

        if re.match(r"^\s*Số\s*:", line_clean, re.IGNORECASE):
            valid_vts = [pb["viet_tat"].upper() for pb in HIDDEN_CONFIG["danh_sach_phong_ban"]]
            notation_match = re.search(r"/\s*([A-ZĐa-zđ0-9]+)(?:\s*-\s*([A-ZĐa-zđ0-9]*))?", line_clean)
            if notation_match:
                full_notation, agency_raw, dept_raw = notation_match.group(0), notation_match.group(1), notation_match.group(2)
                agency_upper, dept = agency_raw.upper(), dept_raw.upper() if dept_raw else ""
                is_notation_err = False
                
                if agency_upper not in MAP_CHUAN_HOA_LOAI:
                    error_list.append(f"❌ Số hiệu sai loại văn bản: `[{agency_raw}]`")
                    is_notation_err = True

                if dept == "":
                    is_notation_err = True
                    correct_agency = MAP_CHUAN_HOA_LOAI.get(agency_upper, agency_raw)
                    choices = [f"Chỉ {correct_agency} (Bỏ dấu gạch ngang)"] + [f"{pb['ten']} ({pb['viet_tat']})" for pb in HIDDEN_CONFIG["danh_sach_phong_ban"]]
                    ambiguous_dict[full_notation] = choices
                    error_list.append(f"⚠️ Số hiệu trống phòng ban: `[{full_notation}]`")
                else:
                    if dept not in valid_vts:
                        is_notation_err = True
                        error_list.append(f"❌ Số hiệu không khớp danh mục BAWACO: `[{full_notation}]`")
                if not is_notation_err:
                    success_items.append(f"Số ký hiệu văn bản chính thức ({full_notation} - Đạt chuẩn)")

    if "hành chánh" in doc_text.lower():
        error_list.append("❌ Sai từ ngữ hành chính: phát hiện từ **hành chánh**.")
        
    if not has_noi_nhan:
        warning_list.append("⚠️ Chưa tìm thấy khối danh mục tiêu đề `Nơi nhận:` (Hệ thống sẽ tự động tạo khi bấm nút Autofix).")

    if not has_nguoi_ky:
        warning_list.append("⚠️ **Cảnh báo:** Chưa phát hiện tiêu đề chức vụ của **Người ký** văn bản ở cuối trang (Ví dụ: GIÁM ĐỐC, TỔNG GIÁM ĐỐC...).")
    else:
        success_items.append("Chức vụ Người ký văn bản (Đã ghi nhận)")

    return success_items, error_list, warning_list, ambiguous_dict

# --- GIAO DIỆN CHÍNH ---
st.title("KIỂM TRA THỂ THỨC VĂN BẢN (NĐ 30)")
st.markdown("🚀 **Phiên bản V6.8 (Trích yếu Fix):** Phân loại định dạng Trích yếu Tờ trình (In đậm) và Công văn (Không in đậm).")
st.markdown("---")

with st.sidebar:
    st.markdown("### 📁 BỘ NẠP VĂN BẢN")
    uploaded_file = st.file_uploader("Chọn file bản thảo văn bản (.docx):", type=["docx"])
    st.markdown("---")
    st.markdown("### 🏢 DANH MỤC BAWACO")
    with st.expander("Xem bảng viết tắt chuẩn", expanded=False):
        st.caption(f"**Công ty:** {HIDDEN_CONFIG['ten_co_quan_me']}")
        for pb in HIDDEN_CONFIG["danh_sach_phong_ban"]:
            st.caption(f"• **{pb['viet_tat']}**: {pb['ten']}")

user_resolutions = {}

if uploaded_file is not None:
    doc = Document(uploaded_file)
    success_items, error_list, warning_list, ambiguous_dict = analyze_document_v6(doc)
    
    st.markdown("### 📊 Kết quả Phân tích & Kiểm tra toàn diện")
    col_overview, col_highlight = st.columns([1, 1])
    
    with col_overview:
        st.write("**📂 Đang xử lý file:** `" + uploaded_file.name + "`")
        with st.expander("Các tiêu chí đã đạt chuẩn", expanded=True):
            for item in success_items: st.write(f"✔️ {item}", unsafe_allow_html=True)
            
    with col_highlight:
        st.markdown("### 🖍️ DANH SÁCH LỖI / SAI LỆCH THỂ THỨC")
        if error_list or warning_list:
            with st.chat_message("assistant", avatar="📝"):
                for item in error_list: st.markdown(item, unsafe_allow_html=True)
                for item in warning_list: st.markdown(item, unsafe_allow_html=True)
        else:
            st.success("✨ Xuất sắc! Văn bản đạt điểm tối đa!")
        
        if ambiguous_dict:
            for notation, choices in ambiguous_dict.items():
                user_resolutions[notation] = st.selectbox(f"Chọn đuôi đúng cho '{notation}':", options=choices)

    # --- HỘP CÔNG CỤ XỬ LÝ (AUTOFIX V6.8) ---
    st.markdown("### 🛠️ HỘP CÔNG CỤ XỬ LÝ TỰ ĐỘNG CHUẨN HOÁ")
    if st.button("🪄 TỰ ĐỘNG FIX TOÀN DIỆN (NĐ30)", type="primary"):
        with st.spinner("Đang định dạng lại lề trang, phông chữ và nội dung..."):
            
            # --- BƯỚC 1: CĂN LỀ TRANG CHUẨN ---
            for section in doc.sections:
                section.top_margin = Cm(2.0)
                section.bottom_margin = Cm(2.0)
                section.left_margin = Cm(3.0)
                section.right_margin = Cm(1.5)

           # --- BƯỚC 2: AUTOFIX THIẾU KHỐI "NƠI NHẬN:" ---
            has_noi_nhan_local = False
            for p in doc.paragraphs:
                if p.text.strip().lower().startswith("nơi nhận"):
                    has_noi_nhan_local = True
                    break
            
            if not has_noi_nhan_local:
                insert_p = None
                for i in range(len(doc.paragraphs)-1, -1, -1):
                    p_text = doc.paragraphs[i].text.strip().lower()
                    if "lưu:" in p_text or "lưu :" in p_text:
                        insert_p = doc.paragraphs[i]
                        for j in range(i-1, max(-1, i-6), -1):
                            prev_text = doc.paragraphs[j].text.strip()
                            if prev_text.startswith("-") or prev_text.startswith("•"):
                                insert_p = doc.paragraphs[j]
                            elif prev_text != "": 
                                break
                        break
                
                if insert_p is not None:
                    new_p = insert_p.insert_paragraph_before()
                    r = new_p.add_run("Nơi nhận:")
                    r.bold = True; r.italic = True; r.font.size = Pt(12); set_font_times(r)

            # --- BƯỚC 3: QUÉT HEADER & ÉP CHUẨN CỠ CHỮ THEO TUYẾN TÍNH ƯU TIÊN ---
            def fix_para(p, paragraph_index):
                text_clean = p.text.replace("|", "").strip()
    
                # BẢO VỆ TUYỆT ĐỐI ĐƯỜNG KẺ NGANG GỐC VÀ DÒNG TRỐNG
                if re.match(r'^[-_=\.\* \t]+$', p.text) or text_clean == "":
                    if re.match(r'^[-_=\.\* \t]+$', p.text):
                        p.alignment = 1
                    return
                                
                # FIX HÀNH CHÁNH -> HÀNH CHÍNH
                if "hành chánh" in p.text.lower():
                    for r in p.runs:
                        if not r.text: continue
                        def fix_case(match):
                            hanh, space, chanh = match.group(1), match.group(2), match.group(3)
                            if chanh.isupper(): chinh = "CHÍNH"
                            elif chanh.istitle(): chinh = "Chính"
                            else: chinh = "chính"
                            return f"{hanh}{space}{chinh}"
                        r.text = re.sub(r"(hành)(\s+)(chánh)", fix_case, r.text, flags=re.IGNORECASE)
                    
                    for i in range(len(p.runs) - 1):
                        r1, r2 = p.runs[i], p.runs[i+1]
                        if r1.text and r2.text:
                            if re.search(r"hành\s+chánh", r1.text + r2.text, flags=re.IGNORECASE):
                                for run_obj in [r1, r2]:
                                    run_obj.text = run_obj.text.replace("chánh", "chính").replace("Chánh", "Chính").replace("CHÁNH", "CHÍNH")

                # Cập nhật lại biến text sau khi đã sửa chính tả "Hành chính"
                text_clean = p.text.replace("|", "").strip()
                text_upper = unicodedata.normalize('NFC', text_clean.upper())
                text_lower = text_clean.lower()

                # ====================================================
                # ƯU TIÊN TUYỆT ĐỐI SỐ 1: KHỐI TRÍCH YẾU NỘI DUNG (Về việc... / V/v...)
                # ====================================================
                if text_lower.startswith("về việc") or text_lower.startswith("v/v"):
                    p.alignment = 1 # Căn giữa chuẩn NĐ 30
                    p.paragraph_format.left_indent = None
                    p.paragraph_format.right_indent = None
                    p.paragraph_format.first_line_indent = None
                    
                    # Khử khoảng trắng thừa đầu câu
                    for r in p.runs:
                        if r.text:
                            r.text = r.text.lstrip("\t").lstrip(" ")
                            break
                            
                    # Phân loại: Công văn (V/v) -> Không đậm. Khác (Về việc) -> In đậm, cỡ 14.
                    is_cong_van = text_lower.startswith("v/v")
                    
                    for r in p.runs:
                        if is_cong_van:
                            r.font.size = Pt(13)
                            r.bold = False
                        else:
                            r.font.size = Pt(14)
                            r.bold = True # Đã ép in đậm lại cho Tờ trình, Quyết định...
                        r.italic = False
                        set_font_times(r)
                    return

                # ====================================================
                # LỚP 1: CÁC TIÊU ĐỀ HEADER NĐ 30
                # ====================================================

                # 1. ÉP CHUẨN QUỐC HIỆU -> Cỡ 12, IN ĐẬM, CHỮ HOA
                if "CỘNG HÒA XÃ HỘI CHỦ NGHĨA VIỆT NAM" in text_upper or re.search(r"CỘNG\s*H[ÒOÀA]+\s*XÃ\s*HỘI", text_upper):
                    p.text = ""
                    r = p.add_run("CỘNG HÒA XÃ HỘI CHỦ NGHĨA VIỆT NAM")
                    r.bold = True; r.font.size = Pt(12); set_font_times(r)
                    p.alignment = 1
                    p.paragraph_format.left_indent = None
                    p.paragraph_format.right_indent = None
                    p.paragraph_format.first_line_indent = None
                    return
                        
                # 2. ÉP CHUẨN TIÊU NGỮ -> Cỡ 13, IN ĐẬM
                if "độc lập" in text_lower and "hạnh phúc" in text_lower:
                    p.text = ""
                    r = p.add_run("Độc lập - Tự do - Hạnh phúc")
                    r.bold = True; r.font.size = Pt(13); set_font_times(r)
                    #rPr = r._r.get_or_add_rPr()
                    #u = OxmlElement('w:u')
                    #u.set(qn('w:val'), 'single') 
                    #u.set(qn('w:sz'), '0.1')
                    #u.set(qn('w:space'), '12')  
                    #rPr.append(u)
                    p.alignment = 1
                    p.paragraph_format.left_indent = None
                    p.paragraph_format.right_indent = None
                    p.paragraph_format.first_line_indent = None
                    return

                # 3. ÉP CHUẨN ĐỊA DANH, NGÀY THÁNG -> Cỡ 14, In nghiêng
                if "ngày" in text_lower and "tháng" in text_lower and "năm" in text_lower:
                    if len(text_clean) < 70 and not any(x in text_lower for x in ["căn cứ", "luật", "nghị định", "quyết định", "thông tư", "v/v", "về việc"]):
                        match_year = re.search(r"năm\s*(\d{4})", text_clean)
                        year_str = match_year.group(1) if match_year else "2026"
                        p.text = f"{HIDDEN_CONFIG['dia_danh_chuan']}, ngày    tháng    năm {year_str}"
                        p.alignment = 1
                        p.paragraph_format.left_indent = None
                        p.paragraph_format.right_indent = None
                        p.paragraph_format.first_line_indent = None
                        for r in p.runs: 
                            r.italic = True; r.font.size = Pt(14); set_font_times(r)
                        return

                # 4. ÉP CHUẨN SỐ KÝ HIỆU VĂN BẢN
                if re.match(r"^\s*Số\s*:", text_clean, re.IGNORECASE) or text_lower.startswith("số:"):
                    temp_text = p.text
                    def fix_notation(match):
                        raw, ag_raw, dt_raw = match.group(0), match.group(1), match.group(2)
                        correct_agency = MAP_CHUAN_HOA_LOAI.get(ag_raw.upper(), ag_raw)
                        valid_vts = [pb["viet_tat"] for pb in HIDDEN_CONFIG["danh_sach_phong_ban"]]
                        valid_vts_upper = [v.upper() for v in valid_vts]
                        try:
                            if raw in user_resolutions:
                                sel = user_resolutions[raw]
                                if "Bỏ dấu gạch ngang" in sel: return f"/{correct_agency}"
                                for pb in HIDDEN_CONFIG["danh_sach_phong_ban"]:
                                    if f"({pb['viet_tat']})" in sel: return f"/{correct_agency}-{pb['viet_tat']}"
                        except NameError: pass
                        dt = dt_raw.upper() if dt_raw else ""
                        if dt in valid_vts_upper: return f"/{correct_agency}-{valid_vts[valid_vts_upper.index(dt)]}"
                        return f"/{correct_agency}-{dt}" if dt else f"/{correct_agency}"

                    temp_text = re.sub(r"^\s*s[ốo]:\s*", "Số: ", temp_text, flags=re.IGNORECASE)
                    temp_text = re.sub(r"/\s*([A-ZĐa-zđ0-9]+)(?:\s*-\s*([A-ZĐa-zđ0-9]*))?", fix_notation, temp_text)
                    p.text = temp_text
                    p.alignment = 1
                    p.paragraph_format.left_indent = None
                    p.paragraph_format.right_indent = None
                    p.paragraph_format.first_line_indent = None
                    for r in p.runs: 
                        r.font.size = Pt(13); r.bold = False; set_font_times(r)
                    return

                # 5. ÉP CHUẨN TÊN CƠ QUAN BAN HÀNH -> Cỡ 13, IN ĐẬM (Khi nằm ở góc trái trang)
                if paragraph_index <= 2:
                    if any(x in text_upper for x in ["CÔNG TY", "CẤP NƯỚC", "PHÒNG", "BAN", "XÍ NGHIỆP", "TRUNG TÂM"]):
                        if not any(x in text_upper for x in ["CỘNG HÒA", "ĐỘC LẬP", "SỐ:", "NGÀY", "THÁNG", "NĂM", "CĂN CỨ", "CHỦ TỊCH", "GIÁM ĐỐC", "BAN QUẢN LÝ", "KÍNH GỬI", "V/V", "VỀ VIỆC"]):
                            p.text = "" 
                            r = p.add_run(text_clean.upper())
                            r.font.size = Pt(13)
                            set_font_times(r)
                            r.bold = True 
                            p.alignment = 1 
                            p.paragraph_format.left_indent = None
                            p.paragraph_format.right_indent = None
                            p.paragraph_format.first_line_indent = None
                            return 

                # 6. ÉP CHUẨN "KÍNH GỬI" -> Cỡ 14
                if text_lower.startswith("kính gửi") or text_lower.startswith("kính gởi"):
                    for r in p.runs:
                        r.bold = False; r.italic = False; r.font.size = Pt(14); set_font_times(r)
                        temp_text = r.text
                        for variant in ["KÍNH GỬI", "Kính Gửi", "KÍNH GỞI", "Kính Gởi"]:
                            temp_text = temp_text.replace(variant, "Kính gửi")
                        r.text = temp_text
                    p.alignment = 0 
                    p.paragraph_format.left_indent = None 
                    p.paragraph_format.first_line_indent = Cm(1.27) 
                    return

                # 7. ÉP CHUẨN TÊN LOẠI VĂN BẢN -> Cỡ 14, IN ĐẬM
                loai_vb_list = ["THÔNG BÁO", "KẾ HOẠCH", "BÁO CÁO", "TỜ TRÌNH", "QUYẾT ĐỊNH", "CHỈ THỊ", "HƯỚNG DẪN", "QUY ĐỊNH", "QUY CHẾ", "PHƯƠNG ÁN"]
                if text_upper in loai_vb_list:
                    for r in p.runs:
                        r.font.size = Pt(14); r.bold = True; set_font_times(r)
                    p.alignment = 1 
                    p.paragraph_format.left_indent = None
                    p.paragraph_format.right_indent = None
                    p.paragraph_format.first_line_indent = None
                    return

                # ====================================================
                # LỚP 2: DÒNG NỘI DUNG CHÍNH (ĐỂ CUỐI CÙNG LÀM BỘ LỌC HẬU BỊ)
                # ====================================================
                tu_khoa_noi_dung = ["vì vậy", "kính mong", "đề nghị", "kính trình", "do đó", "căn cứ", "thực hiện", "nhằm", "để", "sau khi"]
                is_noi_dung = False
                
                if any(text_lower.startswith(tu) for tu in tu_khoa_noi_dung):
                    is_noi_dung = True
                elif len(text_clean) > 20 and not text_clean.isupper():
                    if not text_lower.startswith("kính gửi") and not text_lower.startswith("nơi nhận"):
                        is_noi_dung = True

                if is_noi_dung:
                    p.alignment = 3 
                    for r in p.runs:
                        if r.text:
                            r.text = r.text.lstrip("\t").lstrip(" ")
                            break
                    if not text_clean.startswith("-"):
                        p.paragraph_format.first_line_indent = Cm(1.27)
                        
                    for r in p.runs:
                        r.font.size = Pt(14)
                        set_font_times(r)
                    return

            # --- VÒNG LẶP THỰC THI BƯỚC 3 ---
            for idx, p in enumerate(doc.paragraphs): 
                fix_para(p, idx)
            for t in doc.tables:
                for row in t.rows:
                    for cell in row.cells:
                        for idx, p in enumerate(cell.paragraphs):
                            fix_para(p, idx)

            # --- BƯỚC 4: HẬU XỬ LÝ ĐẶC BIỆT KHỐI TÊN CƠ QUAN Ô TRÁI TABLE ---
            if len(doc.tables) > 0:
                left_cell = doc.tables[0].rows[0].cells[0]
                co_quan_paras = []
                
                for p in left_cell.paragraphs:
                    txt = p.text.strip().lower()
                    if txt.startswith("số:") or re.match(r'^[-_=\.\*]+$', txt):
                        break
                    if txt != "":
                        co_quan_paras.append(p)
                    
                tong_so_dong = len(co_quan_paras)
                for i, p in enumerate(co_quan_paras):
                    p.alignment = 1 
                    p.paragraph_format.left_indent = None
                    p.paragraph_format.right_indent = None
                    p.paragraph_format.first_line_indent = None
                    
                    full_text = re.sub(r'\s+', ' ', p.text.replace("|", "").strip().upper())
                    p.text = ""
                    r = p.add_run(full_text)
                    r.font.size = Pt(13)
                    if tong_so_dong == 1:
                        r.bold = True
                    else:
                        if i == tong_so_dong - 1:
                            r.bold = True
                        else:
                            r.bold = False
                    set_font_times(r)

            # --- BƯỚC 5: POST-PROCESSING KHỐI KÍNH GỬI & NƠI NHẬN ---
            in_noi_nhan_section = False
            in_kinh_gui_section = False 
            
            for p in doc.paragraphs:
                text_clean = p.text.strip()
                text_lower = text_clean.lower()
                
                if re.match(r'^[-_=\.\* \t]+$', p.text):
                    continue

                if text_lower.startswith("kính gửi") or text_lower.startswith("kính gởi"):
                    p.alignment = 0 
                    p.paragraph_format.left_indent = None 
                    p.paragraph_format.first_line_indent = Cm(1.27) 
                    
                    if ":" in text_clean:
                        sau_dau_hai_cham = text_clean.split(":", 1)[1].strip()
                        if len(sau_dau_hai_cham) > 0:
                            in_kinh_gui_section = False 
                            continue
                            
                    in_kinh_gui_section = True
                    continue
                
                if in_kinh_gui_section:
                    if not text_clean: continue
                    if text_lower.startswith("căn cứ") or text_lower.startswith("theo ") or text_lower.startswith("thực hiện") or len(text_clean) > 90:
                        in_kinh_gui_section = False
                    else:
                        p.alignment = 0 
                        p.paragraph_format.left_indent = Cm(3.81) 
                        p.paragraph_format.first_line_indent = None
                        for r in p.runs: 
                            r.bold = False 
                            set_font_times(r)
                        continue

                if len(text_clean) > 0 and (text_lower.startswith("nơi nhận:") or text_lower.startswith("nơi nhận :")):
                    in_noi_nhan_section = True
                    p.text = "" 
                    r = p.add_run("Nơi nhận:")
                    r.bold = True; r.italic = True; r.font.size = Pt(12); set_font_times(r)
                    p.alignment = 0 
                elif in_noi_nhan_section:
                    if text_clean.startswith("-") or text_clean.startswith("+") or text_clean.startswith("•"):
                        p.alignment = 0 
                        p.paragraph_format.left_indent = None 
                        for r in p.runs:
                            r.bold = False; r.italic = False; r.font.size = Pt(11); set_font_times(r)
                    elif "- Lưu:" in text_clean or "- Lưu :" in text_clean:
                        p.alignment = 0 
                        p.paragraph_format.left_indent = None 
                        for r in p.runs:
                            r.bold = False; r.italic = False; r.font.size = Pt(11); set_font_times(r)
                        in_noi_nhan_section = False
                    elif text_clean == "":
                        pass
                    else:
                        in_noi_nhan_section = False
                        
                elif any(x in text_clean.upper() for x in ["GIÁM ĐỐC", "PHÓ GIÁM ĐỐC", "CHỦ TỊCH", "TRƯỞNG PHÒNG", "PHÓ TRƯỞNG PHÒNG", "TỔNG GIÁM ĐỐC"]) and len(text_clean) < 35:
                    if not any(x in text_lower for x in ["căn cứ", "v/v", "về việc", "kính gửi"]):
                        p.alignment = 1 
                        for r in p.runs:
                            r.bold = True; r.italic = False; r.font.size = Pt(14); set_font_times(r)
                            
            # BƯỚC 6: KIỂM TRA LẠI DẤU CÂU (LỖI KHOẢNG TRẮNG TRƯỚC DẤU CHẤM, PHẨY)
            for p in doc.paragraphs:
                if not re.match(r'^[-_=\.\* \t]+$', p.text) and p.text.strip():
                    for r in p.runs:
                        if r.text:
                            r.text = r.text.replace(" ,", ",").replace(" .", ".").replace(" :", ":")
                            
            doc.save("BAWACO_Chuẩn_Hóa_NĐ30.docx")

        with open("BAWACO_Chuẩn_Hóa_NĐ30.docx", "rb") as file:
            st.download_button(
                label="📥 TẢI XUỐNG BẢN ĐÃ FIX (NĐ30)",
                data=file,
                file_name="BAWACO_Chuẩn_Hóa_NĐ30.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                type="primary",
                use_container_width=True
            )
        st.success("✅ Đã hoàn tất sửa lỗi!")
